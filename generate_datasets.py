import argparse
import pandas as pd
import docx
import glob

def calm():
    data = {"user": list(),
            "line": list(),
            "text": list()}
    for file in glob.glob("Raw-Data/Sample_1_CALM/"
                        "Transcripts/Transcripts_PrePostCombined_CreatedByAllisonTackman/"
                        "*.docx"):
        doc = docx.Document(file)
        user  = file.split("/")[-1][:-5]
        data["user"].extend([user for i in range(len(doc.tables[0].rows))])
        data["line"].extend([i for i in range(len(doc.tables[0].rows))])
        data["text"].extend([row.cells[0].text for row in doc.tables[0].rows])

    pd.DataFrame.from_dict(data).to_csv("Data/CALM/CALM.csv")

def breast_cancer():
    data = {"user": list(),
            "line": list(),
            "text": list()}
    for file in glob.glob("Raw-Data/Sample_2_Breast Cancer/Transcripts/"
                          "*.txt"):
        doc = open(file, "rb").readlines()
        user = file.split("/")[-1][:-8]

        data["user"].extend([user for i in range(len(doc))])
        data["line"].extend([i for i in range(len(doc))])
        data["text"].extend([d.decode("utf-8", "ignore").replace("\n", "").replace("\r", "")
                             for d in doc])

    df = pd.DataFrame.from_dict(data)
    df.dropna(subset = ["text"])
    df.to_csv("Data/Breast_Cancer/Breast_Cancer.csv")

def dse():
    data = {"user": list(),
            "line": list(),
            "text": list()}
    for file in glob.glob("Raw-Data/Sample_3_DSE/"
                        "Transcripts/ALL Visits combined/"
                        "*.docx"):
        doc = docx.Document(file)
        user  = file.split("/")[-1].split("_")[0]
        data["user"].extend([user for i in range(len(doc.paragraphs))])
        data["line"].extend([i for i in range(len(doc.paragraphs))])
        data["text"].extend([paragraph.text for paragraph in doc.paragraphs])

    pd.DataFrame.from_dict(data).to_csv("Data/DSE/DSE.csv")

def aging():
    data = {"user": list(),
            "line": list(),
            "text": list()}
    for file in glob.glob("Raw-Data/Sample_4_Aging/"
                          "Transcripts/"
                          "*.docx"):
        doc = docx.Document(file)
        user = file.split("/")[-1][:-5]
        if len(doc.tables) > 0:
            data["user"].extend([user for i in range(len(doc.tables[0].rows))])
            data["line"].extend([i for i in range(len(doc.tables[0].rows))])
            data["text"].extend([row.cells[0].text for row in doc.tables[0].rows])
        else:
            data["user"].extend([user for i in range(len(doc.paragraphs))])
            data["line"].extend([i for i in range(len(doc.paragraphs))])
            data["text"].extend([paragraph.text for paragraph in doc.paragraphs])

    pd.DataFrame.from_dict(data).to_csv("Data/Aging/Aging.csv")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--sample", help="One of the four sample groups, can be "
                                         "CALM, Breast_Cancer, DSE, Aging")

    args = parser.parse_args()

    if args.sample == "CALM":
        calm()
    elif args.sample == "Breast_Cancer":
        breast_cancer()
    elif args.sample == "DSE":
        dse()
    elif args.sample == "Aging":
        aging()
    else:
        print("Wrong sample set selected, please use --sample with correct dataset")
        exit(1)